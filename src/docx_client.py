"""Word document exporter using python-docx."""
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


def add_hyperlink(paragraph, url: str, text: str) -> None:
    """Adds a clickable hyperlink to a paragraph in a .docx document.

    This is a low-level function that manipulates the underlying OOXML to create a
    functional hyperlink.

    Args:
        paragraph: The `docx.paragraph.Paragraph` object to which the link is added.
        url: The URL for the hyperlink.
        text: The display text for the hyperlink.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set hyperlink style (blue and underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


class DocxExporter:
    """Handles the creation and formatting of the .docx Word document.

    This class encapsulates all logic related to `python-docx`, including
    adding formatted text, handling images, and processing hyperlinks.

    Attributes:
        doc: The `docx.Document` object.
        image_position: Where to place images relative to text ('before' or 'after').
        hyperlink_handling: How to treat hyperlinks ('active' or 'ignore').
        include_date_heading: Whether to add date headings for posts.
    """

    def __init__(
        self,
        image_position: str = "after",
        hyperlink_handling: str = "active",
        include_date_heading: bool = True,
    ):
        """Initializes the DocxExporter and creates a new Word document.

        Args:
            image_position: Where to place images relative to text ('before' or 'after').
            hyperlink_handling: How to handle links ('active' or 'ignore').
            include_date_heading: Whether to add a heading for each post's date.
        """
        self.doc = Document()
        self.image_position = image_position
        self.hyperlink_handling = hyperlink_handling
        self.include_date_heading = include_date_heading
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def add_post(self, date: datetime, text: str, images: list[str], entities: list | None = None) -> None:
        """Adds a single post to the Word document.

        This method orchestrates the addition of the date heading, images, and
        formatted text according to the exporter's configuration.

        Args:
            date: The timestamp of the post.
            text: The text content of the post.
            images: A list of local file paths for images to be embedded.
            entities: A list of Telegram MessageEntity objects for formatting.
        """
        if self.include_date_heading:
            date_str = date.strftime("%B %d, %Y")
            self.doc.add_heading(date_str, level=2)

        # Add images before text if configured
        if self.image_position == "before":
            for image_path in images:
                self._process_media_file(image_path)

        # Add text with formatting
        if text:
            self._add_formatted_text(text, entities or [])
            # Add spacing after text
            self.doc.add_paragraph()

        # Add images after text if configured (default)
        if self.image_position == "after":
            for image_path in images:
                self._process_media_file(image_path)

        # Add spacing between posts
        self.doc.add_paragraph()

    def _add_formatted_text(self, text: str, entities: list) -> None:  # noqa: C901
        """Adds a paragraph with text formatted according to Telegram entities.

        Iterates through entities, applying bold, italic, code, and hyperlink
        styles as defined.

        Args:
            text: The raw text of the message.
            entities: A list of `MessageEntity` objects associated with the text.
        """
        if not entities:
            # No formatting, just add plain text
            self.doc.add_paragraph(text)
            return

        # Sort entities by offset
        sorted_entities = sorted(entities, key=lambda e: e.offset)

        # Create paragraph
        paragraph = self.doc.add_paragraph()

        # Track current position in text
        current_pos = 0

        for entity in sorted_entities:
            # Add text before entity
            if entity.offset > current_pos:
                paragraph.add_run(text[current_pos:entity.offset])

            # Get entity text
            entity_text = text[entity.offset:entity.offset + entity.length]

            # Apply formatting based on entity type
            entity_type = type(entity).__name__

            if entity_type == 'MessageEntityBold':
                run = paragraph.add_run(entity_text)
                run.bold = True
            elif entity_type == 'MessageEntityItalic':
                run = paragraph.add_run(entity_text)
                run.italic = True
            elif entity_type == 'MessageEntityCode':
                run = paragraph.add_run(entity_text)
                run.font.name = 'Courier New'
            elif entity_type == 'MessageEntityPre':
                run = paragraph.add_run(entity_text)
                run.font.name = 'Courier New'
            elif entity_type in ('MessageEntityTextUrl', 'MessageEntityUrl'):
                # Add URL as hyperlink if possible
                if entity_type == 'MessageEntityTextUrl':
                    url = getattr(entity, 'url', entity_text)
                else:
                    # Extract URL from text
                    url = entity_text
                    if not url.startswith(('http://', 'https://')):
                        url = 'https://' + url

                if self.hyperlink_handling == "active":
                    # Create actual clickable hyperlink
                    add_hyperlink(paragraph, url, entity_text)
                else:
                    # Just format as blue underlined text (ignore links)
                    run = paragraph.add_run(entity_text)
                    try:
                        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
                    except Exception:
                        # If color setting fails, continue without it
                        pass
                    run.underline = True
            else:
                paragraph.add_run(entity_text)

            current_pos = entity.offset + entity.length

        # Add remaining text after last entity
        if current_pos < len(text):
            paragraph.add_run(text[current_pos:])

    def _process_media_file(self, file_path: str) -> None:
        """Processes a single media file.

        If the file is a supported image, it's embedded in the document. Otherwise,
        a placeholder text is added.

        Args:
            file_path: The local path to the media file.
        """
        SUPPORTED_IMAGE_EXTENSIONS = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
        path = Path(file_path)

        if not path.exists():
            return

        if path.suffix.lower() in SUPPORTED_IMAGE_EXTENSIONS:
            try:
                self._add_image(file_path)
                self.doc.add_paragraph()  # Add spacing after image
            except Exception as e:
                logging.error(f"Error inserting image {file_path}: {e}")
        else:
            # Add a placeholder for unsupported files
            try:
                p = self.doc.add_paragraph()
                run = p.add_run(f"[Unsupported file downloaded: {path.name}]")
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
                self.doc.add_paragraph()  # Add spacing after placeholder
            except Exception as e:
                logging.error(f"Error adding placeholder for {file_path}: {e}")

    def _add_image(self, image_path: str) -> None:
        """Adds an image to the document, scaled to fit the page width.

        Args:
            image_path: The local path to the image file.
        """
        # Get image dimensions
        with Image.open(image_path) as img:
            width_px, height_px = img.size

        # Standard page width in Word is about 6.5 inches (with margins)
        # We'll use 6 inches to account for margins
        max_width = Inches(6)

        # Add image to document
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()

        # Insert image with calculated size
        run.add_picture(image_path, width=max_width)

    def save(self, filename: str) -> None:
        """Saves the Word document to a file.

        Args:
            filename: The path where the document will be saved.
        """
        self.doc.save(filename)
        logging.info(f"Document saved: {filename}")
